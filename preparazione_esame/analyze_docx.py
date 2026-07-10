"""
Analyze 'Appunti Machine learning.docx' — extract structure, images, and formula locations.
This is a read-only research script; it does not modify the original file.
"""

import os
import sys
import io

# Force UTF-8 output to avoid Windows cp1252 encoding errors
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

DOCX_PATH = r"N:\Perceiver_project\preparazione_esame\Appunti Machine learning.docx"
OUT_DIR = r"N:\Perceiver_project\preparazione_esame\temp_images"
os.makedirs(OUT_DIR, exist_ok=True)

doc = Document(DOCX_PATH)

# ── Namespace map for XML queries ──────────────────────────────────────────
nsmap = {
    'w':   'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp':  'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'r':   'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'mc':  'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
    'v':   'urn:schemas-microsoft-com:vml',
    'm':   'http://schemas.openxmlformats.org/officeDocument/2006/math',
    'wp14':'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
    'a14': 'http://schemas.microsoft.com/office/drawing/2010/main',
}


def abbrev(text, maxlen=120):
    text = text.strip()
    if len(text) <= maxlen:
        return text
    return text[:maxlen] + "..."


# ── 1. List all paragraphs ─────────────────────────────────────────────────
print("=" * 80)
print("PARAGRAPH LISTING")
print("=" * 80)
for i, para in enumerate(doc.paragraphs):
    style = para.style.name if para.style else "None"
    text = abbrev(para.text)
    if text:
        print(f"  P{i:04d} [{style:>20s}]  {text}")
    else:
        # Check if the paragraph has images even though no text
        has_img = bool(para._element.findall('.//pic:pic', nsmap) or
                       para._element.findall('.//v:imagedata', nsmap))
        if has_img:
            print(f"  P{i:04d} [{style:>20s}]  <image-only paragraph>")

# ── 2. Detect OMML math blocks ─────────────────────────────────────────────
print("\n" + "=" * 80)
print("OMML MATH BLOCKS (native Word equations)")
print("=" * 80)
math_count = 0
for i, para in enumerate(doc.paragraphs):
    math_elems = para._element.findall('.//m:oMath', nsmap)
    if math_elems:
        math_count += len(math_elems)
        text = abbrev(para.text) if para.text.strip() else "<math-only>"
        print(f"  P{i:04d}  ({len(math_elems)} equation(s))  {text}")
print(f"\nTotal OMML math elements: {math_count}")

# ── 3. Extract every image ─────────────────────────────────────────────────
print("\n" + "=" * 80)
print("IMAGE EXTRACTION")
print("=" * 80)

image_index = 0
rels = doc.part.rels  # relationships at the document-part level

def save_image_blob(blob, para_idx, tag, desc=""):
    """Save raw image bytes and print info."""
    global image_index
    # Try to guess extension from first bytes
    ext = ".png"
    if blob[:3] == b'\xff\xd8\xff':
        ext = ".jpg"
    elif blob[:4] == b'\x89PNG':
        ext = ".png"
    elif blob[:4] == b'GIF8':
        ext = ".gif"
    elif blob[:4] == b'RIFF':
        ext = ".webp"
    elif blob[:2] in (b'BM',):
        ext = ".bmp"
    elif blob[:4] == b'\x00\x00\x01\x00':
        ext = ".ico"
    elif b'<svg' in blob[:200]:
        ext = ".svg"
    else:
        # Check for EMF / WMF
        if blob[:4] == b'\x01\x00\x00\x00':
            ext = ".emf"
        elif blob[:4] == b'\xd7\xcd\xc6\x9a':
            ext = ".wmf"

    fname = f"img_{image_index:03d}_para{para_idx:04d}_{tag}{ext}"
    fpath = os.path.join(OUT_DIR, fname)
    with open(fpath, 'wb') as f:
        f.write(blob)
    size_kb = len(blob) / 1024
    print(f"  [{image_index:3d}] P{para_idx:04d}  {tag:12s}  {size_kb:8.1f} KB  {fname}  {desc}")
    image_index += 1
    return fpath


for i, para in enumerate(doc.paragraphs):
    el = para._element
    context_text = abbrev(para.text, 80)

    # --- Method A: <pic:pic> inside <w:drawing> (modern images) ---
    for pic in el.findall('.//pic:pic', nsmap):
        blip = pic.find('.//a:blip', nsmap)
        if blip is not None:
            r_embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            r_link  = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link')
            rid = r_embed or r_link
            if rid and rid in rels:
                rel = rels[rid]
                try:
                    blob = rel.target_part.blob
                    save_image_blob(blob, i, "drawing", context_text)
                except Exception as exc:
                    print(f"  [ERR] P{i:04d} drawing rId={rid}: {exc}")

    # --- Method B: VML <v:imagedata> (legacy images / equations pasted as images) ---
    for imgdata in el.findall('.//v:imagedata', nsmap):
        rid = imgdata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
        if rid and rid in rels:
            rel = rels[rid]
            try:
                blob = rel.target_part.blob
                save_image_blob(blob, i, "vml", context_text)
            except Exception as exc:
                print(f"  [ERR] P{i:04d} vml rId={rid}: {exc}")

    # --- Method C: OLE / embedded objects ---
    for obj in el.findall('.//w:object', nsmap):
        for imgdata in obj.findall('.//v:imagedata', nsmap):
            rid = imgdata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            if rid and rid in rels:
                rel = rels[rid]
                try:
                    blob = rel.target_part.blob
                    # Avoid duplicates from Method B
                    # (we already handled v:imagedata above, but object-wrapped ones might differ)
                except Exception:
                    pass

# ── 4. Also scan inline shapes via python-docx API ─────────────────────────
print("\n" + "=" * 80)
print("INLINE SHAPES (python-docx API)")
print("=" * 80)
for idx, shape in enumerate(doc.inline_shapes):
    # Find which paragraph this shape belongs to
    shape_elem = shape._inline
    para_elem = shape_elem.getparent().getparent()  # run -> paragraph
    # Find paragraph index
    para_idx = -1
    for pi, p in enumerate(doc.paragraphs):
        if p._element is para_elem:
            para_idx = pi
            break
    w = shape.width
    h = shape.height
    # width/height are in EMU (914400 EMU = 1 inch)
    w_in = w / 914400 if w else 0
    h_in = h / 914400 if h else 0
    print(f"  InlineShape {idx}: type={shape.type}, para={para_idx}, "
          f"size={w_in:.2f}x{h_in:.2f} in")

# ── 5. Summary ──────────────────────────────────────────────────────────────
print("\n" + "=" * 80)
print("SUMMARY")
print("=" * 80)
print(f"  Total paragraphs:      {len(doc.paragraphs)}")
print(f"  Total OMML math blocks:{math_count}")
print(f"  Total images extracted: {image_index}")
print(f"  Total inline shapes:   {len(doc.inline_shapes)}")
print(f"  Output directory:      {OUT_DIR}")

# List saved files
saved = sorted(os.listdir(OUT_DIR))
if saved:
    print(f"\n  Saved files ({len(saved)}):")
    for f in saved:
        fpath = os.path.join(OUT_DIR, f)
        size = os.path.getsize(fpath)
        print(f"    {f}  ({size:,} bytes)")
else:
    print("\n  No image files were saved.")
