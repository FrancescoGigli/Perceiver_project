"""
Replace formula images in a Word document with native OMML equations.

Pipeline: LaTeX -> MathML (via latex2mathml) -> OMML (via MML2OMML.XSL XSLT)
"""

import shutil
import traceback
from docx import Document
from lxml import etree
import latex2mathml.converter

# ── Paths ──────────────────────────────────────────────────────────────────────
SRC = r"N:\Perceiver_project\preparazione_esame\Appunti Machine learning.docx"
DST = r"N:\Perceiver_project\preparazione_esame\Appunti Machine learning - Formule.docx"
XSLT_PATH = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"

# ── Namespaces ─────────────────────────────────────────────────────────────────
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NSMAP = {"w": W_NS, "m": M_NS}

# ── Formula map: (paragraph_index, LaTeX_string) ──────────────────────────────
FORMULA_MAP = [
    (21, r"w_j^{(k+1)} = w_j^{(k)} + \lambda(y_i - \hat{y}_i^{(k)}) x_{ij}"),

    (233, r"\text{Filter} = \begin{bmatrix} 1 & 0 & -1 \\ 1 & 0 & -1 \\ 1 & 0 & -1 \end{bmatrix}"),

    (235, r"\begin{bmatrix} 1 & 0 & -1 \\ 1 & 0 & -1 \\ 1 & 0 & -1 \end{bmatrix} * \begin{bmatrix} P_1 & P_2 & P_3 \\ P_4 & P_5 & P_6 \\ P_7 & P_8 & P_9 \end{bmatrix}"),

    (668, r"I \in \mathbb{R}^{224 \times 224 \times 3}, \quad N = \frac{224}{16} \times \frac{224}{16} = 14 \times 14 = 196"),

    (669, r"p_i \in \mathbb{R}^{768}, \quad x_i = p_i W_E + b_E, \quad W_E \in \mathbb{R}^{768 \times D}, \quad X = [x_1, x_2, \ldots, x_{196}] \in \mathbb{R}^{196 \times 768}"),

    (670, r"Z^0 = [z_{\text{CLS}}, x_1, \ldots, x_{196}] + P, \quad P \in \mathbb{R}^{197 \times 768}, \quad Z^0 \in \mathbb{R}^{197 \times 768}"),

    (673, r"Q = Z^{l-1} W^Q, \; K = Z^{l-1} W^K, \; V = Z^{l-1} W^V, \quad \text{Attn}(Q,K,V) = \text{softmax}\left(\frac{QK^T}{\sqrt{d_k}}\right)V"),

    (674, r"\text{MHA}(Z) = \text{Concat}(\text{head}_1, \ldots, \text{head}_{12}) W^O, \quad \text{FFN}(z) = \text{GELU}(z W_1 + b_1) W_2 + b_2"),

    (677, r"\hat{y} = \text{softmax}(W_{\text{cls}} \, z_{\text{CLS}}^{(12)} + b), \quad W_{\text{cls}} \in \mathbb{R}^{K \times 768}"),

    (678, r"\mathcal{L} = -\sum_{k=1}^{K} y_k \log \hat{y}_k, \quad \frac{\partial \mathcal{L}}{\partial z_{\text{CLS}}^{(12)}} = W_{\text{cls}}^T (\hat{y} - y)"),

    (679, r"\theta \leftarrow \theta - \eta \frac{\partial \mathcal{L}}{\partial \theta}"),

    (970, r"\text{Byte Array} = [r_1, g_1, b_1, r_2, g_2, b_2, \ldots, r_{50176}, g_{50176}, b_{50176}]"),

    (976, r"\gamma(x) = [\cos(2\pi B x), \sin(2\pi B x)] \in \mathbb{R}^{2m}"),

    (990, r"\text{Byte Array} \in \mathbb{R}^{50176 \times 259}"),

    (1050, r"W_k \in \mathbb{R}^{C \times D} = \mathbb{R}^{512 \times 512}"),

    (1051, r"W_v \in \mathbb{R}^{C \times D} = \mathbb{R}^{512 \times 512}"),

    (1054, r"y = W \cdot x + b"),

    (1056, r"W_{ij} \sim U(-a, a), \quad a = \sqrt{\frac{6}{\text{fan}_{\text{in}} + \text{fan}_{\text{out}}}}"),

    (1058, r"W_k \in \mathbb{R}^{259 \times 512}, \quad W_v \in \mathbb{R}^{259 \times 512}"),

    (1060, r"K = \text{ByteArray} \cdot W_k \in \mathbb{R}^{50176 \times 512}, \quad V = \text{ByteArray} \cdot W_v \in \mathbb{R}^{50176 \times 512}"),

    (1101, r"\text{Scores} = Q K^T, \quad Q \in \mathbb{R}^{512 \times 512}, \; K^T \in \mathbb{R}^{512 \times 50176}, \; \text{Scores} \in \mathbb{R}^{512 \times 50176}"),

    (1112, r"\text{Scaled Scores} = \frac{\text{Scores}}{\sqrt{512}} \approx \frac{\text{Scores}}{22.6}"),

    (1121, r"\text{Masked}_{ij} = \begin{cases} \text{Scores}_{ij}, & \text{se Mask}_{ij} = 1 \\ -\infty, & \text{se Mask}_{ij} = 0 \end{cases}"),

    (1137, r"Z_{512 \times 512} = A_{512 \times 50176} \cdot V_{50176 \times 512}"),

    (1146, r"Z' = Z W_o, \quad W_o \in \mathbb{R}^{512 \times 512}"),

    (1151, r"Z_h = A_h V_h, \quad Z = \text{Concat}(Z_1, \ldots, Z_8), \quad Z' = Z W_o \in \mathbb{R}^{512 \times 512}"),

    (1222, r"\text{MLP}(\hat{L}) = \sigma(\hat{L} W_1 + b_1) W_2 + b_2, \quad W_1 \in \mathbb{R}^{512 \times 2048}, \; W_2 \in \mathbb{R}^{2048 \times 512}"),

    (2330, r"X = [x_0; x_1; x_2; x_3] \in \mathbb{R}^{4 \times 192}"),

    (3076, r"Q = I W_Q, \; K = I W_K, \; V = I W_V, \quad A = \text{Softmax}\left(\frac{Q K^T}{\sqrt{D}}\right), \quad Z = A \cdot V"),

    (3084, r"A = \text{Softmax}\left(\frac{Q K^T}{\sqrt{D}}\right), \quad Z = A \cdot V"),

    (3089, r"Q = L W^Q, \; K = X W^K, \; V = X W^V, \quad A = \text{softmax}\left(\frac{Q K^T}{\sqrt{512}}\right), \quad \text{Out} = A V"),

    (3090, r"\text{head}_i = A_i V_i, \quad \text{Out} = \text{Concat}(\text{head}_1, \ldots, \text{head}_8) W^O \in \mathbb{R}^{512 \times 512}"),

    (3094, r"W^O \in \mathbb{R}^{512 \times 512}, \quad d_k = \frac{D}{h} = \frac{512}{8} = 64"),
]


# ── Conversion functions ──────────────────────────────────────────────────────

def load_xslt():
    """Load the MML2OMML XSLT transform once."""
    xslt_tree = etree.parse(XSLT_PATH)
    return etree.XSLT(xslt_tree)


def latex_to_omml(latex_str, xslt_transform):
    """Convert a LaTeX string to an OMML XML element.

    Returns an <m:oMathPara> element wrapping the <m:oMath>.
    """
    # Step 1: LaTeX -> MathML
    mathml_str = latex2mathml.converter.convert(latex_str)

    # Step 2: MathML -> OMML via XSLT
    mathml_tree = etree.fromstring(mathml_str.encode("utf-8"))
    omml_tree = xslt_transform(mathml_tree)
    omml_root = omml_tree.getroot()  # This is <m:oMath>

    # Wrap in <m:oMathPara> for display-style equation
    omath_para = etree.Element(f"{{{M_NS}}}oMathPara")

    # Add centering justification
    omath_para_pr = etree.SubElement(omath_para, f"{{{M_NS}}}oMathParaPr")
    jc = etree.SubElement(omath_para_pr, f"{{{M_NS}}}jc")
    jc.set(f"{{{M_NS}}}val", "center")

    omath_para.append(omml_root)
    return omath_para


def find_runs_with_drawings(paragraph_element):
    """Find all <w:r> elements that contain <w:drawing> children."""
    runs_with_drawings = []
    for run in paragraph_element.findall("w:r", NSMAP):
        drawings = run.findall(".//w:drawing", {"w": W_NS})
        if drawings:
            runs_with_drawings.append(run)
    return runs_with_drawings


def replace_drawings_with_omml(paragraph_element, omml_element):
    """Replace all drawing-containing runs in a paragraph with an OMML element.

    Strategy:
    - Find all <w:r> that contain <w:drawing>
    - Insert the OMML element at the position of the first such run
    - Remove all drawing-containing runs
    """
    runs_with_drawings = find_runs_with_drawings(paragraph_element)
    if not runs_with_drawings:
        return False

    # Insert OMML at the position of the first drawing run
    first_run = runs_with_drawings[0]
    parent = paragraph_element
    parent.insert(list(parent).index(first_run), omml_element)

    # Remove all drawing-containing runs
    for run in runs_with_drawings:
        parent.remove(run)

    return True


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    print(f"Source: {SRC}")
    print(f"Output: {DST}")
    print()

    # Step 1: Copy the file
    shutil.copy2(SRC, DST)
    print("Copied source document to output path.")

    # Step 2: Load the XSLT transform
    xslt_transform = load_xslt()
    print("Loaded MML2OMML.XSL transform.")

    # Step 3: Open the copy
    doc = Document(DST)
    total_paragraphs = len(doc.paragraphs)
    print(f"Document has {total_paragraphs} paragraphs.")
    print()

    # Step 4: Process each formula
    success_count = 0
    fail_count = 0
    skip_count = 0
    failures = []

    for para_idx, latex_str in FORMULA_MAP:
        label = f"Para {para_idx}"
        latex_preview = latex_str[:70] + ("..." if len(latex_str) > 70 else "")

        # Bounds check
        if para_idx >= total_paragraphs:
            msg = f"Paragraph index {para_idx} out of range (max {total_paragraphs - 1})"
            print(f"  SKIP  {label}: {msg}")
            failures.append((para_idx, latex_preview, msg))
            skip_count += 1
            continue

        paragraph = doc.paragraphs[para_idx]
        para_elem = paragraph._element

        # Check that this paragraph actually has drawings
        runs_with_drawings = find_runs_with_drawings(para_elem)
        if not runs_with_drawings:
            msg = "No drawing elements found in this paragraph"
            print(f"  SKIP  {label}: {msg}")
            failures.append((para_idx, latex_preview, msg))
            skip_count += 1
            continue

        # Convert LaTeX -> OMML
        try:
            omml_elem = latex_to_omml(latex_str, xslt_transform)
        except Exception as e:
            msg = f"Conversion error: {e}"
            print(f"  FAIL  {label}: {msg}")
            traceback.print_exc()
            failures.append((para_idx, latex_preview, msg))
            fail_count += 1
            continue

        # Replace drawings with OMML
        try:
            replaced = replace_drawings_with_omml(para_elem, omml_elem)
            if replaced:
                print(f"  OK    {label}: {latex_preview}")
                success_count += 1
            else:
                msg = "replace_drawings_with_omml returned False"
                print(f"  FAIL  {label}: {msg}")
                failures.append((para_idx, latex_preview, msg))
                fail_count += 1
        except Exception as e:
            msg = f"Replacement error: {e}"
            print(f"  FAIL  {label}: {msg}")
            traceback.print_exc()
            failures.append((para_idx, latex_preview, msg))
            fail_count += 1

    # Step 5: Save
    print()
    doc.save(DST)
    print(f"Saved to: {DST}")

    # Step 6: Summary
    print()
    print("=" * 60)
    print("SUMMARY")
    print("=" * 60)
    print(f"  Total formulas:       {len(FORMULA_MAP)}")
    print(f"  Successfully replaced: {success_count}")
    print(f"  Failed conversions:    {fail_count}")
    print(f"  Skipped (no drawing):  {skip_count}")
    print()

    if failures:
        print("FAILURES / SKIPS:")
        for para_idx, latex_preview, msg in failures:
            print(f"  Para {para_idx}: {msg}")
            print(f"    LaTeX: {latex_preview}")
        print()

    print(f"Output file: {DST}")


if __name__ == "__main__":
    main()
